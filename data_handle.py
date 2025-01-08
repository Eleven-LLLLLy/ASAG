# # !/user/bin/env python3
# # -*- coding: utf-8 -*-
# import csv
# import os
# import re
# from itertools import permutations
#
# import numpy as np
# from docx import Document
# import pandas as pd
# def get_score():
#     file_path="/Volumes/Lin11/Data/Evaluation_data/ICNALE_GRA2_1/ICNALE_GRA_2_1.xlsx"
#     sheet_index=5
#     data=pd.read_excel(file_path,sheet_name=sheet_index)
#     if 'Sample Info.1' in data.columns and 'Unnamed: 32' in data.columns:
#         filter_condition = data['Sample Info.1'].str.contains('Essay', na=False)
#         # 筛选符合条件的行，并选择对应列
#         filtered_data = data.loc[filter_condition, ['Sample Info.1', 'Unnamed: 32']]
#         return filtered_data
#     else:
#         print("列 T 或 AG 不存在，请检查数据列名")
# def extract_number(filename):
#     """从文件名中提取数字部分用于排序"""
#     match = re.search(r'(\d{3})', filename)
#     return int(match.group(1)) if match else -1  # 如果没有匹配到，则返回-1
#
# def get_text():
#     folder_path = "/Volumes/Lin11/Data/Evaluation_data/ICNALE_GRA2_1/Rating Samples/ICNALE_GRA_Original_Essays"  # 替换为你的文件夹路径
#
#     # 用于存储所有文件的提取结果
#     all_data = []
#     files = [f for f in os.listdir(folder_path) if f.endswith('.docx') and not f.startswith('.')]
#     sorted_files = sorted(files, key=extract_number)
#     # 遍历文件夹中的所有文件
#     for filename in sorted_files:
#         file_path = os.path.join(folder_path, filename)
#         try:
#             doc = Document(file_path)
#             text=[]
#             for para in doc.paragraphs:
#                 text.append(para.text)
#             all_data.append(text)
#         except Exception as e:
#             print(f"Failed to read {filename}: {e}")
#     return all_data
#
#
# def select_and_pair_scores_texts(scores_df, texts, num_samples=10):
#     """从分数和文本数据中选择一定数量的数据进行成对组合"""
#     scores_array = scores_df['Unnamed: 32'].values  # 转化为NumPy数组
#
#     # 假设scores_df中的索引与texts列表中的顺序一致
#     combined_df = pd.DataFrame({
#         'Score': scores_array,
#         'Text': texts
#     })
#
#     # 按照分数进行等间隔采样，确保分数分布均匀
#     indices = np.round(np.linspace(0, len(combined_df) - 1, num_samples)).astype(int)
#     selected_pairs = combined_df.iloc[indices]
#     print(selected_pairs)
#     return selected_pairs
# def create_all_possible_pairs(dataframe):
#     """创建所有可能的有序对"""
#     items = dataframe.values.tolist()
#     # 使用permutations生成所有可能的有序对
#     pairs = list(permutations(items, 2))
#     return pairs
# def save_pairs_to_csv(pairs, output_file='paired_samples.csv'):
#     """将有序对保存到CSV文件"""
#     with open(output_file, mode='w', newline='', encoding='utf-8') as file:
#         writer = csv.writer(file)
#         writer.writerow(['Score1', 'Text1', 'Score2', 'Text2'])
#         for pair in pairs:
#             writer.writerow([pair[0][0], pair[0][1], pair[1][0], pair[1][1]])
#     print(f"有序对已保存至 {output_file}")
#
# def read_pairs_from_csv(input_file='paired_samples.csv'):
#     """从CSV文件读取有序对"""
#     pairs = []
#     try:
#         df = pd.read_csv(input_file)
#         for _, row in df.iterrows():
#             pairs.append(((row['Score1'], row['Text1']), (row['Score2'], row['Text2'])))
#         print(f"成功从 {input_file} 读取有序对")
#         return pairs
#     except Exception as e:
#         print(f"读取CSV文件时出错: {e}")
#         return []
# if __name__ == "__main__":
#     all_data = get_text()
#     all_score = get_score()
#
#     if all_score is not None and all_data:
#         paired_samples = select_and_pair_scores_texts(all_score, all_data, num_samples=10)
#         all_possible_pairs = create_all_possible_pairs(paired_samples)
#         save_pairs_to_csv(all_possible_pairs)
#         loaded_pairs = read_pairs_from_csv()
#         for idx, pair in enumerate(loaded_pairs[:5], 1):
#             print(f"Pair {idx}:")
#             print(f"First item (Score: {pair[0][0]}, Text: {pair[0][1][:100]}...)")
#             print(f"Second item (Score: {pair[1][0]}, Text: {pair[1][1][:100]}...)")
#             print("-" * 40)
#     else:
#         print("未能获取到有效的分数或文本数据")
# !/user/bin/env python3
# -*- coding: utf-8 -*-
import os
import re
from docx import Document
import pandas as pd
import numpy as np
from itertools import permutations


def get_score():
    file_path = "/Volumes/Lin11/Data/Evaluation_data/ICNALE_GRA2_1/ICNALE_GRA_2_1.xlsx"
    sheet_index = 5
    data = pd.read_excel(file_path, sheet_name=sheet_index)

    if 'Sample Info.1' in data.columns and 'Unnamed: 32' in data.columns:
        filter_condition = data['Sample Info.1'].str.contains('Essay', na=False)
        # 筛选符合条件的行，并选择对应列
        filtered_data = data.loc[filter_condition, ['Sample Info.1', 'Unnamed: 32']]
        return filtered_data
    else:
        print("列 Sample Info.1 或 Unnamed: 32 不存在，请检查数据列名")
        return None


def extract_number(filename):
    """从文件名中提取数字部分用于排序"""
    match = re.search(r'(\d{3})', filename)
    return int(match.group(1)) if match else -1  # 如果没有匹配到，则返回-1


def get_text():
    folder_path = "/Volumes/Lin11/Data/Evaluation_data/ICNALE_GRA2_1/Rating Samples/ICNALE_GRA_Original_Essays"  # 替换为你的文件夹路径

    all_data = []
    files = [f for f in os.listdir(folder_path) if f.endswith('.docx') and not f.startswith('.')]
    sorted_files = sorted(files, key=extract_number)

    for filename in sorted_files:
        file_path = os.path.join(folder_path, filename)
        try:
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])  # 合并所有段落为一个字符串
            all_data.append(text)
        except Exception as e:
            print(f"Failed to read {filename}: {e}")
    return all_data


def select_and_pair_scores_texts(scores_df, texts, num_samples=10):
    """从分数和文本数据中选择一定数量的数据进行成对组合"""
    scores_array = scores_df['Unnamed: 32'].values  # 转化为NumPy数组

    # 假设scores_df中的索引与texts列表中的顺序一致
    combined_df = pd.DataFrame({
        'Score': scores_array,
        'Text': texts
    })

    # 按照分数进行等间隔采样，确保分数分布均匀
    indices = np.round(np.linspace(0, len(combined_df) - 1, num_samples)).astype(int)
    selected_pairs = combined_df.iloc[indices]

    return selected_pairs


def create_all_possible_pairs(dataframe):
    """创建所有可能的有序对"""
    items = dataframe.values.tolist()
    # 使用permutations生成所有可能的有序对
    pairs = list(permutations(items, 2))
    return pairs


def format_and_save_answers(pairs, output_file='formatted_answers.txt'):
    """根据给定的格式整理文本并保存到文件"""
    with open(output_file, mode='w', encoding='utf-8') as file:
        for idx, pair in enumerate(pairs, 1):
            file.write(f"answer 1:\n")
            file.write(f'"{pair[0][1]}"\n')  # 第一个文本
            file.write(f"answer 2:\n")
            file.write(f'"{pair[1][1]}"\n')  # 第二个文本
            file.write("-" * 40 + "\n\n")
    print(f"格式化的答案已保存至 {output_file}")


if __name__ == "__main__":
    all_data = get_text()
    all_score = get_score()

    if all_score is not None and all_data:
        paired_samples = select_and_pair_scores_texts(all_score, all_data, num_samples=10)
        all_possible_pairs = create_all_possible_pairs(paired_samples)

        # 格式化并保存文本答案
        format_and_save_answers(all_possible_pairs)

        # 示例：读取刚刚保存的答案文件
        with open('formatted_answers.txt', 'r', encoding='utf-8') as file:
            print(file.read())
    else:
        print("未能获取到有效的分数或文本数据")